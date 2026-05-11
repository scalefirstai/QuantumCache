"""
DOCX parser — AFME DDQs (custodian, CSD, prime broker).

Uses python-docx. Heading hierarchy from style names (Heading 1 … Heading 4)
defines section boundaries. Sections become SectionAnchor.item; sub-headings
become SectionAnchor.subsection. AFME DDQs use numbered question identifiers
(e.g., "1.1", "B.3.4") that we capture as the section_id leaf.
"""

from __future__ import annotations

import datetime as dt
import io
import re
from typing import Iterable

import docx

from .span import (
    Provenance,
    SectionAnchor,
    Span,
    make_span,
)

PARSER_VERSION = "python-docx@1"

MAX_SPAN_CHARS = 2000
MIN_SPAN_CHARS = 20

QUESTION_ID_RE = re.compile(r"^\s*([A-Z]?\d+(?:\.\d+){0,4})\s+")


def _slug(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_")
    return s[:80] or "Section"


def _heading_level(paragraph) -> int | None:
    """Return 1-based heading level if the paragraph is a Heading style."""
    style = (paragraph.style.name or "").strip()
    m = re.match(r"Heading\s+(\d+)", style)
    if not m:
        return None
    return int(m.group(1))


def _split_long_text(text: str) -> list[str]:
    if len(text) <= MAX_SPAN_CHARS:
        return [text]
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z(])", text)
    out: list[str] = []
    buf = ""
    for s in sentences:
        if len(buf) + len(s) + 1 > MAX_SPAN_CHARS and buf:
            out.append(buf.strip())
            buf = s
        else:
            buf = (buf + " " + s).strip() if buf else s
    if buf.strip():
        out.append(buf.strip())
    return out


def _emit_text(
    text: str, headings: list[str], doc_id: str, doc_hash: str,
    provenance, ordinal_state: list[int], yield_to: list,
) -> None:
    """Helper: emit one or more spans for a piece of text, attached to current heading hierarchy."""
    if len(text) < MIN_SPAN_CHARS:
        return
    m = QUESTION_ID_RE.match(text)
    subsection = m.group(1) if m else None
    item = " > ".join(headings) if headings else "Document"
    section_id = _slug(item) + (f".q{subsection}" if subsection else "")
    for chunk in _split_long_text(text):
        anchor = SectionAnchor(item=item, subsection=subsection, doc_hash=doc_hash)
        yield_to.append(make_span(
            doc_id=doc_id,
            doc_hash=doc_hash,
            section_id=section_id,
            ordinal=ordinal_state[0],
            text=chunk,
            anchor=anchor,
            provenance=provenance,
        ))
        ordinal_state[0] += 1


def parse_docx(
    *,
    body: bytes,
    doc_id: str,
    doc_hash: str,
    source: str,
    extra: dict | None = None,
) -> Iterable[Span]:
    """Parse DOCX preserving paragraph + table content in document order.

    AFME DDQs hold most question text inside tables; iterating only paragraphs
    misses ~95% of content. We walk `document.element.body` in XML order and
    handle <w:p> and <w:tbl> nodes as we hit them.
    """
    document = docx.Document(io.BytesIO(body))
    provenance = Provenance(
        source=source,
        parser=PARSER_VERSION,
        ingested_at=dt.datetime.now(dt.timezone.utc).isoformat(),
        extra=extra or {},
    )
    headings: list[str] = ["Front"]
    ordinal_state = [0]   # mutable counter
    out: list[Span] = []

    body_el = document.element.body
    # Build a lookup from element -> Paragraph / Table for fast access.
    para_by_el = {p._element: p for p in document.paragraphs}
    table_by_el = {t._element: t for t in document.tables}

    for child in body_el.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            para = para_by_el.get(child)
            if para is None:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            level = _heading_level(para)
            if level is not None:
                headings = headings[: max(0, level - 1)]
                headings.append(text)
                ordinal_state[0] = 0
                continue
            _emit_text(text, headings, doc_id, doc_hash, provenance, ordinal_state, out)
        elif tag == "tbl":
            table = table_by_el.get(child)
            if table is None:
                continue
            for row in table.rows:
                # Concatenate non-empty cells with " | " — keeps Q + A pairing.
                cells = [c.text.strip() for c in row.cells]
                cells = [c for c in cells if c]
                if not cells:
                    continue
                # Deduplicate horizontally-merged cells (python-docx repeats merged cell text).
                deduped: list[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                text = " | ".join(deduped)
                _emit_text(text, headings, doc_id, doc_hash, provenance, ordinal_state, out)
    yield from out
